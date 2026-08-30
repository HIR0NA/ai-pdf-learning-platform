from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from PIL import Image

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "outputs" / "demo-screenshots-report"
OUT.mkdir(parents=True, exist_ok=True)
DOCX = OUT / "ai-study-companion-demo-evidence.docx"

BLUE = "2E74B5"
INK = "1F2937"
MUTED = "5B6573"
LIGHT = "F2F4F7"
GREEN = "166534"
AMBER = "92400E"

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tcMar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tcMar.append(node)
        node.set(qn('w:w'), str(v)); node.set(qn('w:type'), 'dxa')

def set_table_widths(table, widths):
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW'); tblPr.append(tblW)
    tblW.set(qn('w:w'), str(sum(widths))); tblW.set(qn('w:type'), 'dxa')
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for w in widths:
        col = OxmlElement('w:gridCol'); col.set(qn('w:w'), str(w)); grid.append(col)
    for row in table.rows:
        for cell, w in zip(row.cells, widths):
            cell.width = Inches(w / 1440)
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is None: tcW = OxmlElement('w:tcW'); tcPr.append(tcW)
            tcW.set(qn('w:w'), str(w)); tcW.set(qn('w:type'), 'dxa')
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def font(run, size=11, color=INK, bold=False, italic=False):
    run.font.name = 'Calibri'
    run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), 'Calibri')
    run._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'), 'Calibri')
    run.font.size = Pt(size); run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold; run.italic = italic

def add_para(doc, text='', size=11, color=INK, bold=False, italic=False, align=None, before=0, after=6):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after); p.paragraph_format.line_spacing = 1.1
    if align is not None: p.alignment = align
    if text:
        r = p.add_run(text); font(r, size, color, bold, italic)
    return p

def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f'Heading {level}')
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text); font(r, {1:16,2:13,3:12}[level], BLUE if level < 3 else '1F4D78', True)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after = Pt(4); p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text); font(r, 10.5, INK); return p

def add_caption(doc, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text); font(r, 9, MUTED, italic=True); return p

def add_screenshot(doc, filename, title, description, note=None):
    heading(doc, title, 2)
    add_para(doc, description, 10.5, MUTED, after=6)
    path = ROOT / 'screenshots' / filename
    with Image.open(path) as im:
        w, h = im.size
    max_w = 6.25
    max_h = 7.0
    width = max_w
    height = width * h / w
    if height > max_h:
        height = max_h; width = height * w / h
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(path), width=Inches(width), height=Inches(height))
    add_caption(doc, f"หลักฐานภาพหน้าจอ: {filename}")
    if note:
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(10)
        r = p.add_run('หมายเหตุ: '); font(r, 9.5, AMBER, True)
        r = p.add_run(note); font(r, 9.5, MUTED)

doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.8); sec.bottom_margin = Inches(0.75); sec.left_margin = Inches(0.85); sec.right_margin = Inches(0.85)
sec.header_distance = Inches(0.35); sec.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles['Normal']; normal.font.name = 'Calibri'; normal.font.size = Pt(11); normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.1
for lvl, size, color, before, after in [(1,16,BLUE,16,8),(2,13,BLUE,12,6),(3,12,'1F4D78',8,4)]:
    st = styles[f'Heading {lvl}']; st.font.name='Calibri'; st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(color)
    st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after); st.paragraph_format.keep_with_next=True

# Running header/footer
hp = sec.header.paragraphs[0]; hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = hp.add_run('AI Study Companion  |  Demo Evidence'); font(r, 8.5, MUTED)
fp = sec.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = fp.add_run('จัดทำเพื่อการตรวจสอบหน้าจอและสิทธิ์การเข้าถึง  •  25 สิงหาคม 2026'); font(r, 8.5, MUTED)

# Title block (memo masthead)
add_para(doc, 'DEMO EVIDENCE REPORT', 10, BLUE, True, after=4)
add_para(doc, 'AI Study Companion', 24, '000000', True, after=3)
add_para(doc, 'รายงานหลักฐานหน้าจอ: Login, Student Dashboard, Admin/Database และ RBAC', 13, MUTED, after=14)
for label, value in [('โครงการ', 'AI PDF Learning Platform'), ('ขอบเขต', 'ตรวจสอบภาพหน้าจอการใช้งานและการแยกสิทธิ์ Student/Admin'), ('วันที่จัดทำ', '25 สิงหาคม 2026')]:
    p = doc.add_paragraph(); p.paragraph_format.space_after=Pt(2)
    r=p.add_run(label+': '); font(r,10.5,INK,True); r=p.add_run(value); font(r,10.5,INK)

add_para(doc, 'สรุปผล', 12, BLUE, True, before=14, after=5)
add_para(doc, 'เอกสารฉบับนี้รวบรวมภาพหน้าจอจากสภาพแวดล้อมทดสอบของระบบ เพื่อใช้เป็นหลักฐานประกอบการนำเสนอและการทำ Visual QA โดยแยกมุมมองผู้ใช้ Student ออกจาก Admin และบันทึกข้อจำกัดของการทดสอบ RBAC ที่เกิดจากเบราว์เซอร์บล็อกหน้า 403 โดยตรง', 10.5, INK, after=8)

table = doc.add_table(rows=1, cols=3); table.alignment=WD_TABLE_ALIGNMENT.LEFT; set_table_widths(table,[900,3600,4860])
hdr = table.rows[0].cells
for c, t in zip(hdr, ['#','หลักฐาน','สถานะ']):
    set_cell_shading(c, LIGHT); p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.LEFT; p.paragraph_format.space_after=Pt(0); r=p.add_run(t); font(r,9.5,INK,True)
for i,(e, label, s) in enumerate([('01','หน้า Login','เก็บภาพแล้ว'),('02','Admin Dashboard / Database evidence','เก็บภาพแล้ว'),('03','Student Dashboard','เก็บภาพแล้ว'),('04','RBAC / 403 evidence','เก็บภาพแล้ว พร้อมหมายเหตุ')],1):
    cells=table.add_row().cells
    for c in cells: set_cell_margins(c)
    for c,t in zip(cells,[str(i),label,s]):
        p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0); r=p.add_run(t); font(r,9.5,INK)
    if i < 4: set_cell_shading(cells[2], 'ECFDF5')
add_para(doc, 'ตารางสรุปหลักฐานที่แนบในรายงาน', 9, MUTED, italic=True, after=10, align=WD_ALIGN_PARAGRAPH.CENTER)

add_screenshot(doc, '01-login.png', '1. หน้า Login', 'แสดงจุดเริ่มต้นของการเข้าสู่ระบบและการเลือกใช้บัญชีตาม Role', 'ภาพนี้เป็นหลักฐานการแสดงผลของหน้า Login ในสภาพแวดล้อมทดสอบ ไม่แสดงรหัสผ่านหรือ Secret ใด ๆ')
doc.add_page_break()
add_screenshot(doc, '03-dashboard.png', '2. Student Dashboard', 'แสดงมุมมองฝั่งนักศึกษา: ภาพรวมเอกสาร การสนทนากับ AI เวลาเรียน และพื้นที่อัปโหลดเอกสาร', 'เมนูและป้ายกำกับ STUDENT ช่วยยืนยันว่าหน้านี้เป็นมุมมองผู้ใช้ทั่วไป และไม่มีเมนู Admin Panel')
doc.add_page_break()
add_screenshot(doc, '02-admin-database.png', '3. Admin Dashboard / Database Evidence', 'แสดงมุมมองผู้ดูแลระบบพร้อมตัวเลขสรุป ผู้ใช้ เอกสาร ข้อความ และ Security Logs รวมถึงรายการผู้ใช้และ Role', 'ข้อมูลในภาพเป็นข้อมูลจากฐานข้อมูลทดสอบ ณ เวลาที่เก็บภาพ ควรใช้เพื่อสาธิตโครงสร้างการควบคุมและการตรวจสอบ ไม่ใช่ข้อมูล Production')
doc.add_page_break()
add_screenshot(doc, '04-rbac-403.png', '4. RBAC Test: 403 Forbidden', 'แสดงหลักฐานประกอบการทดสอบการปฏิเสธการเข้าถึงเส้นทางที่สงวนไว้สำหรับ Admin', 'เบราว์เซอร์แสดงหน้า “localhost is blocked / ERR_BLOCKED_BY_CLIENT” เมื่อเปิดเส้นทาง 403 โดยตรง จึงไม่ใช่ภาพข้อความ JSON 403 จากแอปโดยตรง อย่างไรก็ตาม server-side guard ใน src/proxy.ts และ src/app/api/admin/overview/route.ts เป็นจุดบังคับสิทธิ์ที่ควรอ้างอิงร่วมกับภาพนี้')

heading(doc, 'ข้อเสนอแนะสำหรับการเก็บหลักฐานรอบถัดไป', 1)
add_bullet(doc, 'บันทึกผล API test ของบัญชี Student ที่เรียก /api/admin/overview ให้เห็น status 403 และ response body โดยตรง (เช่นจาก DevTools หรือ automated test)')
add_bullet(doc, 'เก็บภาพ Admin Panel ที่มีเมนูเฉพาะ Admin และภาพ Student ที่ไม่มีเมนูดังกล่าวใน viewport เดียวกัน เพื่อเปรียบเทียบ Role ได้ชัดเจน')
add_bullet(doc, 'ก่อนเผยแพร่รายงาน ให้ตรวจว่าไม่มี Token, API Key, อีเมลจริง หรือข้อมูลส่วนบุคคลในภาพหน้าจอ')

doc.save(DOCX)
print(DOCX)
